# # backend/app/utils.py
# import json, os, re
# from typing import Any, Dict, List, Optional
# from jinja2 import Environment, FileSystemLoader, select_autoescape
# from sumy.parsers.plaintext import PlaintextParser
# from sumy.nlp.tokenizers import Tokenizer
# from sumy.summarizers.lsa import LsaSummarizer
# from sumy.nlp.stemmers import Stemmer
# from sumy.utils import get_stop_words
# from youtube_transcript_api import YouTubeTranscriptApi
# from duckduckgo_search import ddg
# from io import BytesIO
# from PIL import Image
# import requests
# import uuid
# import base64
# from app import config
# from fastapi import HTTPException

# BASE_DIR = os.path.dirname(__file__)
# TEMPLATES_DIR = os.path.join(BASE_DIR, "templates")
# os.makedirs(config.OUTPUT_PPT_DIR, exist_ok=True)
# os.makedirs(config.OUTPUT_IMAGE_DIR, exist_ok=True)

# env = Environment(loader=FileSystemLoader(TEMPLATES_DIR), autoescape=select_autoescape(["txt"]))

# # -------------------------
# # Template utilities
# # -------------------------
# def render_template(name: str, **context) -> str:
#     tpl = env.get_template(name)
#     return tpl.render(**context)

# # -------------------------
# # JSON coercion (robust)
# # -------------------------
# def coerce_json(text: str) -> Any:
#     if not text or not isinstance(text, str):
#         raise ValueError("No text to parse as JSON")
#     try:
#         return json.loads(text)
#     except Exception:
#         pass
#     # find largest JSON block
#     candidates = []
#     stack = []
#     current_start = None
#     for i, ch in enumerate(text):
#         if ch in ['{', '[']:
#             if current_start is None:
#                 current_start = i
#             stack.append(ch)
#         elif ch in ['}', ']'] and stack:
#             stack.pop()
#             if not stack and current_start is not None:
#                 candidates.append(text[current_start:i+1])
#                 current_start = None
#     candidates.sort(key=len, reverse=True)
#     for c in candidates:
#         try:
#             return json.loads(c)
#         except Exception:
#             try:
#                 return json.loads(c.replace("'", '"'))
#             except Exception:
#                 continue
#     # last try: replace single quotes
#     try:
#         return json.loads(text.replace("'", '"'))
#     except Exception as e:
#         raise ValueError(f"Failed to coerce JSON: {str(e)}")

# # -------------------------
# # SUMY extractive summarizer wrapper
# # -------------------------
# def sumy_summarize(text: str, sentences_count: int = 3, language: str = "english") -> str:
#     if not text or len(text.strip()) == 0:
#         return ""
#     parser = PlaintextParser.from_string(text, Tokenizer(language))
#     stemmer = Stemmer(language)
#     summarizer = LsaSummarizer(stemmer)
#     summarizer.stop_words = get_stop_words(language)
#     sentences = summarizer(parser.document, sentences_count)
#     return " ".join(str(s) for s in sentences)

# # -------------------------
# # DuckDuckGo search wrapper (ddg)
# # -------------------------
# def regulated_ddg_search(query: str, depth: int = 3) -> str:
#     if not config.WEB_SEARCH_ALLOWED or not query:
#         return ""
#     try:
#         results = ddg(query, region="wt-wt", safesearch="Off", time="y", max_results=5)
#         snippets = []
#         for r in results[:depth]:
#             snippet = r.get("body") or r.get("desc") or r.get("title") or ""
#             snippet = re.sub(r"\s+", " ", snippet).strip()
#             snippets.append(snippet)
#         return "\n".join(snippets)
#     except Exception:
#         return ""

# # -------------------------
# # YouTube transcript extraction
# # -------------------------
# def fetch_youtube_transcript(video_url_or_id: str) -> str:
#     try:
#         video_id = video_url_or_id.split("v=")[-1].split("&")[0]
#         parts = video_id.split("/")  # if user passed id directly or short URL
#         if len(parts) > 1:
#             video_id = parts[-1]
#         transcript_list = YouTubeTranscriptApi.get_transcript(video_id)
#         return " ".join([t["text"] for t in transcript_list])
#     except Exception as e:
#         return ""

# def fetch_and_summarize_youtube(links: List[str], depth: int = 2, summ_sentences: int = 2) -> str:
#     if not config.YOUTUBE_ALLOWED or not links:
#         return ""
#     summaries = []
#     for link in links[:depth]:
#         text = fetch_youtube_transcript(link)
#         if text:
#             summary = sumy_summarize(text, sentences_count=summ_sentences)
#             summaries.append(summary)
#     return "\n".join(summaries)

# # -------------------------
# # File text extraction (txt/pdf/docx)
# # -------------------------
# def extract_text_from_file_bytes(file_bytes: bytes, ext: str) -> str:
#     ext = ext.lower()
#     if ext == "txt":
#         return file_bytes.decode("utf-8", errors="ignore")
#     elif ext == "pdf":
#         try:
#             from io import BytesIO
#             from PyPDF2 import PdfReader
#             reader = PdfReader(BytesIO(file_bytes))
#             pages = []
#             for p in reader.pages:
#                 pages.append(p.extract_text() or "")
#             return "\n".join(pages)
#         except Exception:
#             return ""
#     elif ext in ("docx", "doc"):
#         try:
#             import docx
#             from io import BytesIO
#             doc = docx.Document(BytesIO(file_bytes))
#             return "\n".join([p.text for p in doc.paragraphs])
#         except Exception:
#             return ""
#     else:
#         return ""

# # -------------------------
# # Image helpers (save url or base64)
# # -------------------------
# def save_image_from_url_or_b64(url_or_b64: str, out_dir: Optional[str] = None) -> Optional[str]:
#     out_dir = out_dir or config.OUTPUT_IMAGE_DIR
#     os.makedirs(out_dir, exist_ok=True)
#     try:
#         if url_or_b64.startswith("http"):
#             r = requests.get(url_or_b64, timeout=30)
#             r.raise_for_status()
#             img = Image.open(BytesIO(r.content)).convert("RGBA")
#         else:
#             # assume base64
#             if "," in url_or_b64:
#                 payload = url_or_b64.split(",", 1)[1]
#             else:
#                 payload = url_or_b64
#             b = base64.b64decode(payload)
#             img = Image.open(BytesIO(b)).convert("RGBA")
#         out_path = os.path.join(out_dir, f"{uuid.uuid4().hex}.png")
#         img.save(out_path)
#         return out_path
#     except Exception:
#         return None

#=====================================================================
import os
import io
import json
import logging
import requests
from typing import Dict, Any, Optional

# File parsing
from PyPDF2 import PdfReader
from docx import Document

# Summarization & NLP
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lsa import LsaSummarizer
from jinja2 import Template

import nltk
nltk.download("punkt")

# Search & YouTube
from duckduckgo_search import DDGS
from youtube_transcript_api import YouTubeTranscriptApi

# ============================================================
# CONFIG
# ============================================================
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "")
OPENROUTER_BASE_URL = "https://openrouter.ai/api/v1"

logger = logging.getLogger(__name__)

# app/utils.py
from string import Template
import os

def render_template(template_name: str, **kwargs):
    """
    Simple template renderer using string.Template.
    template_name: name of template file in templates/ folder
    kwargs: key-value pairs for template replacement
    """
    template_path = os.path.join(os.path.dirname(__file__), "templates", template_name)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")
    
    with open(template_path, "r", encoding="utf-8") as f:
        content = f.read()
    
    # Use safe_substitute to avoid KeyErrors if some keys are missing
    return Template(content).safe_substitute(**kwargs)


# ============================================================
# FILE HELPERS
# ============================================================
def extract_text_from_file_bytes(file_bytes: bytes, filename: str) -> str:
    """
    Extract text from uploaded file bytes. Supports PDF and DOCX.
    """
    ext = filename.lower().split(".")[-1]
    text = ""

    if ext == "pdf":
        pdf_reader = PdfReader(io.BytesIO(file_bytes))
        for page in pdf_reader.pages:
            text += page.extract_text() or ""
    elif ext in ["docx", "doc"]:
        doc = Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            text += para.text + "\n"
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    return text.strip()


# ============================================================
# GENERIC CALL FUNCTION
# ============================================================
def call_openrouter_model(
    model: str,
    prompt: str,
    negative_prompt: Optional[str] = None,
    size: str = "1024x1024",
    n: int = 1,
) -> Optional[Dict[str, Any]]:
    """
    Call OpenRouter API for a given model (Stable Diffusion, LLaMA-4-Maverick, etc.).
    Returns JSON response or None if failed.
    """
    url = f"{OPENROUTER_BASE_URL}/images"
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
    }

    payload = {
        "model": model,
        "prompt": prompt,
        "size": size,
        "n": n,
    }

    if negative_prompt:
        payload["negative_prompt"] = negative_prompt

    try:
        resp = requests.post(url, headers=headers, json=payload, timeout=60)
        if resp.status_code == 200:
            return resp.json()
        else:
            logger.error(f"[OpenRouter:{model}] Error {resp.status_code}: {resp.text}")
            return None
    except Exception as e:
        logger.error(f"[OpenRouter:{model}] Exception: {str(e)}")
        return None


# ============================================================
# PRIMARY IMAGE GEN: Stable Diffusion (Standard)
# ============================================================
def call_primary_image_gen(
    prompt: str,
    negative_prompt: Optional[str] = None,
    size: str = "1024x1024",
    n: int = 1,
) -> Optional[Dict[str, Any]]:
    """
    Calls Stable Diffusion (Standard) via OpenRouter for image generation.
    """
    logger.info("[ImageGen] Using primary: Stable Diffusion (standard)")
    return call_openrouter_model(
        model="stabilityai/stable-diffusion-2-1",
        prompt=prompt,
        negative_prompt=negative_prompt,
        size=size,
        n=n,
    )


# ============================================================
# FALLBACK IMAGE GEN: LLaMA-4-Maverick
# ============================================================
def call_fallback_image_gen(
    prompt: str,
    negative_prompt: Optional[str] = None,
    size: str = "1024x1024",
    n: int = 1,
) -> Optional[Dict[str, Any]]:
    """
    Calls LLaMA-4-Maverick (SD-based captioning/alt gen) via OpenRouter as fallback.
    """
    logger.warning("[ImageGen] Primary failed, using fallback: LLaMA-4-Maverick")
    return call_openrouter_model(
        model="meta-llama/llama-4-maverick",
        prompt=prompt,
        negative_prompt=negative_prompt,
        size=size,
        n=n,
    )


# ============================================================
# IMAGE HELPER
# ============================================================
def extract_image_urls(response: Dict[str, Any]) -> list:
    """
    Extracts image URLs from OpenRouter response.
    """
    if not response or "data" not in response:
        return []
    return [item.get("url") for item in response.get("data", []) if "url" in item]


# ============================================================
# TEMPLATE RENDERING
# ============================================================
def render_template(template_str: str, context: dict) -> str:
    """Render a Jinja2 template string with provided context."""
    try:
        template = Template(template_str)
        return template.render(**context)
    except Exception as e:
        logger.error(f"Template rendering failed: {e}")
        return template_str


# ============================================================
# SUMMARIZATION
# ============================================================
def sumy_summarize(text: str, sentence_count: int = 5) -> str:
    """Summarize text using Sumy LSA summarizer."""
    try:
        parser = PlaintextParser.from_string(text, Tokenizer("english"))
        summarizer = LsaSummarizer()
        summary = summarizer(parser.document, sentence_count)
        return " ".join(str(sentence) for sentence in summary)
    except Exception as e:
        logger.error(f"Summarization failed: {e}")
        return text


# ============================================================
# DUCKDUCKGO SEARCH
# ============================================================
# def regulated_ddg_search(query: str, max_results: int = 5) -> list:
#     """Perform DuckDuckGo search and return list of dicts with title + href."""
#     try:
#         results = []
#         with DDGS() as ddgs:
#             for r in ddgs.text(query, max_results=max_results):
#                 results.append({
#                     "title": r.get("title"),
#                     "href": r.get("href")
#                 })
#         return results
#     except Exception as e:
#         logger.error(f"DDG search failed: {e}")
#         return []
def regulated_ddg_search(query: str, max_results: int = 5, depth: int = None) -> list:
    # Use depth if provided; otherwise fallback to max_results
    results_to_fetch = depth if depth is not None else max_results
    try:
        results = []
        with DDGS() as ddgs:
            for r in ddgs.text(query, max_results=results_to_fetch):
                results.append({
                    "title": r.get("title"),
                    "href": r.get("href")
                })
        return results
    except Exception as e:
        logger.error(f"DDG search failed: {e}")
        return []


# ============================================================
# YOUTUBE TRANSCRIPT + SUMMARY
# ===========================================================
import logging
logger = logging.getLogger(__name__)

# backend/app/utils.py

def fetch_and_summarize_youtube(links: list, depth: int = 1, sentence_count: int = 2):
    """
    Fetch transcripts from the first `depth` YouTube links and summarize them safely.
    Returns a single string combining all summaries.
    """
    summaries = []

    # Only take the first `depth` links
    links_to_process = links[:depth]

    for link in links_to_process:
        try:
            from youtube_transcript_api import YouTubeTranscriptApi
            video_id = link.split("v=")[-1].split("&")[0]  # Extract video ID
            transcript_list = YouTubeTranscriptApi.get_transcript(video_id)
            full_text = " ".join([t.get("text", "") for t in transcript_list])

            # Use your sumy_summarize function
            from app.utils import sumy_summarize
            summary = sumy_summarize(full_text, sentence_count=sentence_count)
            summaries.append(summary)

        except Exception as e:
            # Instead of breaking, include a safe message
            summaries.append(f"(Failed to fetch {link}: {str(e)})")

    return "\n".join(summaries)


# ============================================================
# JSON HELPER
# ============================================================
def coerce_json(data) -> dict:
    """Safely coerce input to JSON."""
    try:
        if isinstance(data, str):
            return json.loads(data)
        elif isinstance(data, dict):
            return data
        else:
            return {}
    except Exception as e:
        logger.error(f"JSON coercion failed: {e}")
        return {}

def extract_image_urls(api_response: dict) -> list:
    """
    Extract image URLs from Stable Diffusion/OpenRouter API response
    """
    try:
        urls = []
        if "data" in api_response:
            for item in api_response["data"]:
                if "url" in item:
                    urls.append(item["url"])
        return urls
    except Exception as e:
        logger.error(f"[ImageGen] Failed to extract image URLs: {e}")
        return []
